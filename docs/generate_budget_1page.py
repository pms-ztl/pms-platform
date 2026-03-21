"""
PMS AI Budget Report — 1 Page, Simple, USD + INR
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# Narrow margins for max space
for section in doc.sections:
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(0.8)
    section.left_margin = Cm(1.3)
    section.right_margin = Cm(1.3)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(9)
style.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
style.paragraph_format.space_after = Pt(2)
style.paragraph_format.space_before = Pt(0)

VIOLET = RGBColor(0x6D, 0x28, 0xD9)
BLUE = RGBColor(0x1E, 0x3A, 0x5F)
GRAY = RGBColor(0x64, 0x74, 0x8B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK = "1E293B"
ALT = "F1F5F9"

def shade(cell, color):
    tc = cell._element.get_or_add_tcPr()
    shd = tc.makeelement(qn('w:shd'), {qn('w:fill'): color, qn('w:val'): 'clear'})
    tc.append(shd)

def make_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0]
    for i, txt in enumerate(headers):
        c = hdr.cells[i]
        shade(c, DARK)
        p = c.paragraphs[0]
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.space_before = Pt(1)
        r = p.add_run(txt)
        r.bold = True
        r.font.size = Pt(8)
        r.font.color.rgb = WHITE
        r.font.name = 'Calibri'
    for ri, row_data in enumerate(rows):
        row = t.add_row()
        for ci, txt in enumerate(row_data):
            c = row.cells[ci]
            if ri % 2 == 1:
                shade(c, ALT)
            p = c.paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(1)
            r = p.add_run(str(txt))
            r.font.size = Pt(8)
            r.font.name = 'Calibri'
            if ci == 0:
                r.bold = True
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    return t

# ── TITLE ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(0)
r = p.add_run('PMS Platform \u2014 Agentic AI Budget Report')
r.font.size = Pt(16)
r.font.color.rgb = VIOLET
r.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(6)
r = p.add_run('70 AI Agents  |  5 LLM Providers  |  64 Tools  |  March 2026  |  Confidential')
r.font.size = Pt(8)
r.font.color.rgb = GRAY

# ── SECTION 1: MODEL PRICING ──
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(3)
r = p.add_run('MODEL PRICING  ')
r.font.size = Pt(10)
r.font.color.rgb = BLUE
r.bold = True
r = p.add_run('(1 USD = \u20b991.92)')
r.font.size = Pt(8)
r.font.color.rgb = GRAY

make_table(
    ['Model', 'Provider', 'Input / 1K Tokens', 'Output / 1K Tokens', 'Quality', 'Best For'],
    [
        ['Claude Sonnet 4', 'Anthropic', '$0.003 (\u20b90.28)', '$0.015 (\u20b91.38)', 'Excellent', 'Production'],
        ['GPT-4o', 'OpenAI', '$0.0025 (\u20b90.23)', '$0.010 (\u20b90.92)', 'Excellent', 'Fallback'],
        ['GPT-4o Mini', 'OpenAI', '$0.00015 (\u20b90.01)', '$0.0006 (\u20b90.06)', 'Good', 'Testing'],
        ['Gemini 2.0 Flash', 'Google', '$0.0001 (\u20b90.01)', '$0.0004 (\u20b90.04)', 'Good', 'Economy'],
        ['DeepSeek Chat', 'DeepSeek', '$0.00014 (\u20b90.01)', '$0.00028 (\u20b90.03)', 'Good', 'Budget'],
    ],
    [1.3, 0.8, 1.3, 1.3, 0.7, 1.0]
)

# ── SECTION 2: COST PER INTERACTION ──
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(3)
r = p.add_run('COST PER INTERACTION  ')
r.font.size = Pt(10)
r.font.color.rgb = BLUE
r.bold = True
r = p.add_run('(using Claude Sonnet 4)')
r.font.size = Pt(8)
r.font.color.rgb = GRAY

make_table(
    ['Type', 'LLM Calls', 'Tokens Used', 'Cost (USD)', 'Cost (INR)'],
    [
        ['Simple Chat', '1', '2K \u2013 5K', '$0.01 \u2013 $0.03', '\u20b91 \u2013 \u20b93'],
        ['Agentic Task', '3 \u2013 5', '15K \u2013 30K', '$0.05 \u2013 $0.20', '\u20b95 \u2013 \u20b918'],
        ['Help Assistant', '1', '2K \u2013 3K', '$0.005 \u2013 $0.02', '\u20b90.5 \u2013 \u20b92'],
        ['System Scan', '1 \u2013 2', '2K \u2013 5K', '$0.01 \u2013 $0.05', '\u20b91 \u2013 \u20b95'],
    ],
    [1.2, 0.8, 1.2, 1.3, 1.3]
)

# ── SECTION 3: MONTHLY COST ──
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(3)
r = p.add_run('MONTHLY RUNNING COST  ')
r.font.size = Pt(10)
r.font.color.rgb = BLUE
r.bold = True
r = p.add_run('(recurring \u2014 paid every month while users actively use AI)')
r.font.size = Pt(8)
r.font.color.rgb = GRAY

make_table(
    ['Users', 'All-Premium* (USD)', 'All-Premium* (INR)', 'Smart Routing** (USD)', 'Smart Routing** (INR)'],
    [
        ['50', '$220 \u2013 $780', '\u20b920,222 \u2013 \u20b971,698', '$90 \u2013 $320', '\u20b98,273 \u2013 \u20b929,414'],
        ['100', '$440 \u2013 $1,550', '\u20b940,445 \u2013 \u20b91,42,476', '$180 \u2013 $640', '\u20b916,546 \u2013 \u20b958,829'],
        ['250', '$1,100 \u2013 $3,870', '\u20b91,01,112 \u2013 \u20b93,55,526', '$450 \u2013 $1,600', '\u20b941,364 \u2013 \u20b91,47,072'],
        ['500', '$2,200 \u2013 $7,750', '\u20b92,02,224 \u2013 \u20b97,12,380', '$900 \u2013 $3,200', '\u20b982,728 \u2013 \u20b92,94,144'],
        ['1,000', '$4,400 \u2013 $15,450', '\u20b94,04,448 \u2013 \u20b914,20,116', '$1,800 \u2013 $6,400', '\u20b91,65,456 \u2013 \u20b95,88,288'],
    ],
    [0.6, 1.3, 1.5, 1.2, 1.5]
)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(1)
r = p.add_run('*All-Premium = Every AI query uses Claude Sonnet 4 (most expensive, highest quality).\n')
r.font.size = Pt(7.5)
r.font.color.rgb = GRAY
r.italic = True
r = p.add_run('**Smart Routing = Simple questions use cheap Gemini Flash; only complex tasks use Claude Sonnet 4. Same features, 55% cheaper.')
r.font.size = Pt(7.5)
r.font.color.rgb = GRAY
r.italic = True

# ── SECTION 4: DECISION MATRIX ──
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(3)
r = p.add_run('QUICK DECISION  ')
r.font.size = Pt(10)
r.font.color.rgb = BLUE
r.bold = True
r = p.add_run('(for 100 users)')
r.font.size = Pt(8)
r.font.color.rgb = GRAY

make_table(
    ['Scenario', 'Model', 'USD / Month', 'INR / Month', 'API Key Needed'],
    [
        ['Demo / Testing', 'Gemini 2.0 Flash', '$8 \u2013 $25', '\u20b9735 \u2013 \u20b92,298', 'GEMINI_API_KEY'],
        ['Pilot', 'GPT-4o Mini', '$15 \u2013 $50', '\u20b91,379 \u2013 \u20b94,596', 'OPENAI_API_KEY'],
        ['Production', 'Claude Sonnet 4', '$440 \u2013 $1,550', '\u20b940,445 \u2013 \u20b91,42,476', 'ANTHROPIC_API_KEY'],
        ['Production (Best Value)', 'Claude + Gemini', '$180 \u2013 $640', '\u20b916,546 \u2013 \u20b958,829', 'Both keys'],
    ],
    [1.3, 1.3, 1.1, 1.5, 1.3]
)

# ── SAFEGUARDS (compact) ──
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(3)
r = p.add_run('BUILT-IN SAFEGUARDS  ')
r.font.size = Pt(10)
r.font.color.rgb = BLUE
r.bold = True

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
r = p.add_run(
    '\u2022 Max 50,000 tokens/task ($0.50 cap)   '
    '\u2022 Rate limit: 15 calls/user/min   '
    '\u2022 Monthly budget: configurable (default $100/tenant)   '
    '\u2022 Redis cache: 1hr TTL   '
    '\u2022 Circuit breaker: auto-fallback across 5 providers   '
    '\u2022 Human approval for risky actions'
)
r.font.size = Pt(8)

# ── SECTION 6: FULL TESTING BUDGET ──
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(3)
r = p.add_run('ONE-TIME TESTING BUDGET  ')
r.font.size = Pt(10)
r.font.color.rgb = BLUE
r.bold = True
r = p.add_run('(paid once \u2014 to test all 70 agents + orchestration before go-live)')
r.font.size = Pt(8)
r.font.color.rgb = GRAY

make_table(
    ['Testing Phase', 'What', 'Input Tokens', 'Output Tokens'],
    [
        ['Smoke Test', '70 agents \u00d7 2 calls each', '700K', '200K'],
        ['Functional Test', '70 agents \u00d7 3 agentic tasks', '5.25M', '1.5M'],
        ['Orchestration', '30 multi-agent scenarios', '1.5M', '500K'],
        ['RBAC & Security', '5 roles \u00d7 20 test cases', '500K', '150K'],
        ['Edge Cases', '50 error/recovery scenarios', '1M', '300K'],
        ['Load Simulation', 'Concurrent users, rate limits', '2M', '600K'],
        ['Regression (2 rounds)', 'Re-run key tests', '3M', '900K'],
        ['Buffer', 'Debugging, retries', '1M', '350K'],
        ['TOTAL', '', '~15M', '~4.5M'],
    ],
    [1.3, 2.2, 1.1, 1.1]
)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(3)
r = p.add_run('ONE-TIME TESTING COST  ')
r.font.size = Pt(10)
r.font.color.rgb = BLUE
r.bold = True
r = p.add_run('(~19.5M tokens total \u2014 this is NOT per user, it is the total to test the whole platform once)')
r.font.size = Pt(8)
r.font.color.rgb = GRAY

make_table(
    ['Model', 'Input Cost', 'Output Cost', 'One-Time Total (USD)', 'One-Time Total (INR)'],
    [
        ['Gemini 2.0 Flash', '$1.50', '$1.80', '$3 \u2013 $5', '\u20b9276 \u2013 \u20b9460'],
        ['GPT-4o Mini', '$2.25', '$2.70', '$5 \u2013 $8', '\u20b9460 \u2013 \u20b9735'],
        ['Claude Sonnet 4', '$45', '$67.50', '$100 \u2013 $150', '\u20b99,192 \u2013 \u20b913,788'],
        ['Smart Mix (Best)', 'Gemini bulk + Claude critical', '', '$25 \u2013 $35', '\u20b92,298 \u2013 \u20b93,217'],
    ],
    [1.3, 1.1, 0.9, 1.3, 1.5]
)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(1)
r = p.add_run('Smart Mix = Use cheap Gemini Flash to test most agents, use expensive Claude Sonnet 4 only for testing complex multi-step tasks. Best of both worlds.')
r.font.size = Pt(7.5)
r.font.color.rgb = GRAY
r.italic = True

# ── BIG CONCLUSION ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(10)
p.paragraph_format.space_after = Pt(2)
r = p.add_run('RECOMMENDATION')
r.font.size = Pt(9)
r.font.color.rgb = GRAY
r.bold = True

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(1)
r = p2.add_run('One-Time Testing (all 70 agents): $25\u2013$35 (\u20b92,298\u2013\u20b93,217)')
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x14, 0xB8, 0xA6)  # teal
r.bold = True

p2a = doc.add_paragraph()
p2a.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2a.paragraph_format.space_after = Pt(4)
r = p2a.add_run('Buy ~20M tokens from Google AI Studio (Gemini Flash) + small Anthropic credit for Claude validation')
r.font.size = Pt(8)
r.font.color.rgb = GRAY
r.italic = True

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_after = Pt(1)
r = p3.add_run('Monthly Production (100 users): $180\u2013$640/mo (\u20b916,546\u2013\u20b958,829/mo)')
r.font.size = Pt(14)
r.font.color.rgb = VIOLET
r.bold = True

p3a = doc.add_paragraph()
p3a.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3a.paragraph_format.space_after = Pt(4)
r = p3a.add_run('Buy Anthropic API credits (Claude Sonnet 4) + Google AI Studio credits (Gemini Flash for economy)')
r.font.size = Pt(8)
r.font.color.rgb = GRAY
r.italic = True

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
p4.paragraph_format.space_after = Pt(2)
r = p4.add_run('Smart model routing saves 55%  |  No upfront license fee \u2014 pay only for tokens you use')
r.font.size = Pt(9)
r.font.color.rgb = BLUE
r.bold = True

# Save
output_path = os.path.join(os.path.dirname(__file__), 'PMS_AI_Budget_v5.docx')
doc.save(output_path)
print(f'Saved: {output_path}')
