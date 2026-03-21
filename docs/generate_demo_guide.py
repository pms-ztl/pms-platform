"""
Generate PMS Platform Demo Preparation Guide (.docx)
Last-minute non-technical guide for investors, vendors & project managers.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Styles ──────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)  # dark navy
    h.font.name = 'Calibri'

# Helper
def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(f' — {text}')
    else:
        p.add_run(text)
    return p

def add_table_row(table, cells_data, bold_first=False):
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(9.5)
        if bold_first and i == 0:
            run.bold = True

def shade_cells(row, color):
    for cell in row.cells:
        shading = cell._element.get_or_add_tcPr()
        shd = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): color,
            qn('w:val'): 'clear',
        })
        shading.append(shd)


# ═══════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('PMS Platform')
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x6D, 0x28, 0xD9)  # violet
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Demo Preparation Guide')
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

doc.add_paragraph()
tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = tagline.add_run('Non-Technical Quick Reference for Investors, Vendors & Project Managers')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
run.italic = True

doc.add_paragraph()
version = doc.add_paragraph()
version.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = version.add_run('Version 1.0  |  March 2026  |  Confidential')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ═══════════════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. What is PMS Platform? (30-Second Elevator Pitch)',
    '2. The 12 Unique Selling Points (USPs)',
    '3. How the App Works — Feature Walkthrough',
    '4. The Brain Behind It — CPIS Scoring System',
    '5. How Math & AI Power the Scores',
    '6. Agentic AI — Your 70 Virtual HR Assistants',
    '7. Who Sees What — Roles & Access Control (RBAC)',
    '8. Super Admin vs Tenant Admin',
    '9. Multi-Tenant SaaS Architecture (Simple)',
    '10. Subscription Plans & Feature Gating',
    '11. Real-Time & Live Updates',
    '12. Security & Compliance',
    '13. Key Numbers to Remember',
    '14. Demo Flow — Recommended 5-Minute Script',
    '15. 30 FAQs with Quick Answers',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 1. ELEVATOR PITCH
# ═══════════════════════════════════════════════════════════
doc.add_heading('1. What is PMS Platform?', level=1)
doc.add_heading('The 30-Second Elevator Pitch', level=3)

doc.add_paragraph(
    'PMS (Performance Management System) is an enterprise-grade, AI-powered SaaS platform '
    'that helps organizations manage, measure, and improve employee performance. '
    'Think of it as the "brain" of HR — it replaces spreadsheets, manual reviews, '
    'and guesswork with data-driven intelligence.'
)

doc.add_paragraph()
doc.add_heading('In Simple Words:', level=3)
add_bullet('It tracks what employees are working on (Goals & OKRs)')
add_bullet('It measures how well they are doing (Performance Scores)')
add_bullet('It gives fair, unbiased ratings (AI-powered Calibration)')
add_bullet('It helps managers make better decisions (Analytics & Insights)')
add_bullet('It grows employee careers (Development Plans & Skill Maps)')
add_bullet('It keeps the workplace healthy (Engagement & Culture Tracking)')
add_bullet('It does all of this across multiple organizations securely (Multi-Tenant SaaS)')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('One-liner: ')
run.bold = True
p.add_run('"PMS replaces your entire HR performance stack with one intelligent platform powered by 70 AI agents."')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 2. USPs
# ═══════════════════════════════════════════════════════════
doc.add_heading('2. The 12 Unique Selling Points (USPs)', level=1)

usps = [
    ('CPIS Score — The Performance DNA',
     'A single 0-100 score that captures 8 dimensions of performance using 17+ statistical formulas. '
     'No other platform has this level of mathematical rigor. It shows letter grades (A+ to F), '
     'star ratings, confidence intervals, and growth trajectory.'),
    ('70 Specialized AI Agents (Neural Swarm)',
     'Not just one chatbot — 70 specialized AI agents organized in 6 clusters: '
     'Bio-Performance, Hyper-Learning, Liquid Workforce, Culture & Empathy, Governance, and Core. '
     'Each agent is an expert in its domain.'),
    ('Agentic AI — AI That Takes Action',
     'Unlike basic AI chatbots that only answer questions, our AI agents can actually DO things — '
     'draft reviews, create goals, schedule meetings, run analytics, and generate reports autonomously.'),
    ('Built-in Bias Detection & Fairness',
     'The system automatically detects reviewer bias, self-assessment inflation, and rating manipulation. '
     'Uses the 4/5ths rule (disparate impact analysis) used by legal courts to ensure fairness.'),
    ('Real-Time Performance Tracking',
     'Performance scores update in real-time via WebSocket connections. Managers see live dashboards '
     'with activity heatmaps, workload distribution, and anomaly detection.'),
    ('Multi-Tenant Architecture',
     'One platform serves unlimited organizations, each fully isolated. Data from Company A '
     'can never be seen by Company B. Each tenant has its own subscription, settings, and users.'),
    ('360-Degree Review System',
     'Feedback from managers, peers, direct reports, self-assessment, and external reviewers — '
     'all calibrated to remove bias and weighted by reviewer trust scores.'),
    ('Engagement & Culture Analytics',
     'eNPS scores, pulse surveys, mood tracking, burnout risk detection, and department-level '
     'culture diagnostics — all in one place.'),
    ('Performance Improvement Plans (PIP)',
     'Structured workflows for underperforming employees with milestones, check-ins, and outcome tracking. '
     'Legally defensible documentation.'),
    ('Calibration Sessions',
     'Statistical tools for HR to normalize ratings across managers and departments. '
     'Detects outliers, suggests discussion points, and records adjustment decisions.'),
    ('Enterprise Security',
     'Role-based access control (RBAC) with 5 role levels, cross-tenant access blocking, '
     'brute force detection, audit logging, and delegation of authority.'),
    ('Progressive Web App (PWA)',
     'Works on any device — desktop, tablet, mobile — without installing from app stores. '
     'Offline capability, push notifications, and native-like experience.'),
]

for i, (title_text, desc) in enumerate(usps, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'USP #{i}: {title_text}')
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x6D, 0x28, 0xD9)
    doc.add_paragraph(desc)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 3. FEATURE WALKTHROUGH
# ═══════════════════════════════════════════════════════════
doc.add_heading('3. How the App Works — Feature Walkthrough', level=1)
doc.add_paragraph(
    'The platform has 130+ features organized into 11 categories. '
    'Here are the key ones you need to know:'
)

categories = [
    ('Core Performance Management', [
        ('Goals & OKRs', 'Employees set goals linked to company objectives. Progress is tracked with completion %, quality ratings, and timeliness.'),
        ('Performance Reviews', 'Structured review cycles (quarterly/annual). Supports self, manager, peer, and 360-degree reviews.'),
        ('Feedback', 'Continuous feedback — praise, constructive, suggestions. Includes sentiment analysis and skill tagging.'),
        ('Evidence Collection', 'Employees attach proof of achievements. Impact and quality scored. Linked to reviews and promotions.'),
    ]),
    ('Analytics & Intelligence', [
        ('CPIS Dashboard', 'The main dashboard shows a single score (0-100) with 8 dimension breakdown, letter grade, stars, and trend.'),
        ('29 Interactive Charts', 'Performance timeline, skill gap radar, goal progress funnel, feedback sentiment, engagement heatmap, and more.'),
        ('Anomaly Detection', 'AI automatically flags unusual patterns (sudden drops, rating manipulation, gaming).'),
        ('Leaderboards', 'Performance rankings by team, department, or organization — daily, weekly, monthly.'),
    ]),
    ('People Development', [
        ('Skill Gap Analysis', 'Visual radar chart showing current skills vs target skills with gap identification.'),
        ('Development Plans', 'Personalized growth plans with learning activities, milestones, and progress tracking.'),
        ('Career Pathing', 'AI-suggested career trajectories based on skills, performance, and organizational needs.'),
        ('Mentoring Hub', 'Mentor/mentee matching based on skills and goals.'),
    ]),
    ('Engagement & Culture', [
        ('Pulse Surveys', 'Quick daily/weekly mood and energy check-ins with anonymous trend tracking.'),
        ('eNPS Score', 'Employee Net Promoter Score — measures how likely employees recommend the company.'),
        ('1-on-1 Meetings', 'Scheduling, agenda, notes, and action item tracking for manager-employee meetings.'),
        ('Recognition', 'Peer-to-peer recognition with points, badges, and company-wide visibility.'),
    ]),
    ('Decision Support', [
        ('Calibration', 'HR-led sessions to normalize ratings across teams and remove manager bias.'),
        ('Compensation', 'Data-driven pay decisions linked to performance evidence and market benchmarks.'),
        ('Promotion', 'AI-scored promotion readiness with evidence-backed recommendations.'),
        ('PIP Management', 'Performance Improvement Plans with structured milestones and legal documentation.'),
    ]),
    ('Communication', [
        ('Real-time Chat', 'Direct messages, group channels, team conversations with emoji reactions and threading.'),
        ('Announcements', 'Company-wide or targeted announcements with read tracking.'),
        ('Notifications', 'Real-time push notifications via WebSocket + email digest.'),
    ]),
]

for cat_name, features in categories:
    doc.add_heading(cat_name, level=2)
    for feat_name, feat_desc in features:
        add_bullet(feat_desc, bold_prefix=feat_name)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 4. CPIS SCORING
# ═══════════════════════════════════════════════════════════
doc.add_heading('4. The Brain Behind It — CPIS Scoring System', level=1)

doc.add_paragraph(
    'CPIS stands for Comprehensive Performance Intelligence Score. '
    'It is a single number (0 to 100) that represents the overall performance of an employee. '
    'Think of it like a credit score — but for work performance.'
)

doc.add_heading('The 8 Dimensions (What Makes Up the Score)', level=2)

dims_table = doc.add_table(rows=1, cols=4)
dims_table.style = 'Light Grid Accent 1'
dims_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = dims_table.rows[0].cells
for i, text in enumerate(['#', 'Dimension', 'Weight', 'What It Measures']):
    hdr[i].text = text

dims = [
    ('1', 'Goal Attainment', '25%', 'How well are goals being completed? Are they on time? High quality?'),
    ('2', 'Review Quality', '20%', 'What do managers & peers say in reviews? (bias-corrected)'),
    ('3', 'Feedback Sentiment', '12%', 'Is the feedback received mostly positive or negative?'),
    ('4', 'Collaboration Impact', '10%', 'Does the person help others? Cross-team work? Feedback given?'),
    ('5', 'Consistency', '10%', 'Reliable delivery? On-time? Steady velocity? No wild swings?'),
    ('6', 'Growth Trajectory', '8%', 'Are scores improving over time? New skills? Training completed?'),
    ('7', 'Evidence Quality', '8%', 'Quality and diversity of submitted work evidence'),
    ('8', 'Initiative & Innovation', '7%', 'Going above and beyond? Mentoring? Process improvements?'),
]
for d in dims:
    add_table_row(dims_table, d, bold_first=True)

doc.add_paragraph()
doc.add_heading('The Output', level=2)

outputs = [
    ('Score', '0-100 (e.g., 78.5)'),
    ('Letter Grade', 'A+, A, B+, B, C+, C, D, F'),
    ('Star Rating', '1 to 5 stars'),
    ('Rank Label', 'E.g., "High Achiever", "Strong Contributor", "Exceptional Performer"'),
    ('Confidence', 'How reliable is this score? (narrow or wide margin based on data volume)'),
    ('Growth Direction', 'Improving, Stable, or Declining — based on trend analysis'),
]
for label, desc in outputs:
    add_bullet(desc, bold_prefix=label)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 5. MATH & AI
# ═══════════════════════════════════════════════════════════
doc.add_heading('5. How Math & AI Power the Scores', level=1)

doc.add_paragraph(
    'The platform uses 17+ statistical formulas. Here\'s what each does in plain English:'
)

math_table = doc.add_table(rows=1, cols=3)
math_table.style = 'Light Grid Accent 1'
hdr = math_table.rows[0].cells
for i, text in enumerate(['Formula', 'Plain English', 'Where Used']):
    hdr[i].text = text

maths = [
    ('Weighted Average', 'Some things matter more than others. Multiply by importance, then average.', 'Goal scores, dimension weights'),
    ('Harmonic Mean', 'A fairer average for ratings — prevents one high score from hiding many low ones.', 'Review ratings'),
    ('EWMA (Exponential Smoothing)', 'Recent data matters more than old data. Like a "memory" that fades.', 'Feedback trends, quality scores'),
    ('Z-Score', 'How far above or below average someone is. Like grading on a curve.', 'Team comparisons, anomaly detection'),
    ('Sigmoid Function', 'Turns any number into a smooth 0-to-100 scale. Small inputs barely count, then rapid growth, then plateaus.', 'Collaboration scores, risk scores'),
    ('Standard Deviation', 'How spread out scores are. Low = consistent team. High = big gaps.', 'Team analytics, calibration'),
    ('Linear Regression', 'Drawing a trend line through scores. Positive slope = improving.', 'Growth trajectory, predictions'),
    ('Bayesian Estimation', 'When data is limited, borrow from department averages. More data = trust individual more.', 'New employee fairness'),
    ('Shannon Entropy', 'How evenly distributed ratings are. 0 = everyone got the same rating. 1 = perfectly spread.', 'Rating distribution health'),
    ('Gini Coefficient', 'Inequality measure (like wealth inequality). Are top performers getting ALL the recognition?', 'Team equity analysis'),
    ('Pearson Correlation', 'Do two things move together? E.g., does attendance correlate with performance?', 'Analytics insights'),
    ('Percentile Rank', 'What % of people scored lower? "82nd percentile" = better than 82% of peers.', 'Rankings, leaderboards'),
    ('Disparate Impact (4/5ths Rule)', 'Legal fairness test. If one group scores below 80% of the top group, flag bias.', 'Bias detection'),
    ('Calibrated Ratings', 'Removes reviewer bias by normalizing each reviewer\'s rating habits.', 'Review fairness'),
    ('Risk Score', 'Probability of goal failure based on pace, complexity, and dependencies.', 'Goal risk assessment'),
]
for m in maths:
    add_table_row(math_table, m)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Key takeaway for investors: ')
run.bold = True
p.add_run('Every number in the system is mathematically derived — zero guesswork, zero hardcoded values. '
          'Missing data defaults to neutral (50/100), not penalizing. The system is auditable and legally defensible.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 6. AGENTIC AI
# ═══════════════════════════════════════════════════════════
doc.add_heading('6. Agentic AI — Your 70 Virtual HR Assistants', level=1)

doc.add_heading('What Is Agentic AI?', level=2)
doc.add_paragraph(
    'Regular AI chatbots answer questions. Agentic AI takes action. '
    'When an employee says "draft my self-review," the AI doesn\'t just suggest what to write — '
    'it actually reads past goals, feedback, and achievements, then writes a complete review draft.'
)

doc.add_heading('How It Works (Simple)', level=2)
steps = [
    'User sends a message (e.g., "Help me prepare for my review")',
    'System classifies intent: Is this a simple question or a multi-step task?',
    'If multi-step: The Agentic Engine creates a plan with steps',
    'Each step is routed to the best specialist agent (e.g., Performance Agent, Goal Agent)',
    'Agents execute independently — fetching data, running calculations, creating content',
    'For risky actions (like submitting something), it asks the user for approval first',
    'Results are combined and presented back to the user',
    'Everything is logged in an audit trail',
]
for i, step in enumerate(steps, 1):
    p = doc.add_paragraph(f'{i}. {step}')

doc.add_heading('The 6 Agent Clusters', level=2)

clusters = [
    ('Core Agents (20)', 'Performance tracking, review drafting, goal intelligence, coaching, career planning, NLP queries, compensation analysis, workforce insights, compliance, reporting'),
    ('Bio-Performance (10)', 'Burnout prevention, stress monitoring, focus optimization, sleep analysis, break optimization, ergonomics, energy management — monitors human wellbeing signals'),
    ('Hyper-Learning (12)', 'Skill gap forecasting, micro-learning delivery, knowledge brokering, career simulation, cross-training, certification tracking — accelerates skill development'),
    ('Liquid Workforce (10)', 'Internal gig marketplace, salary benchmarking, equity tracking, retirement planning, vendor negotiations, succession planning — manages flexible talent'),
    ('Culture & Empathy (10)', 'Culture analytics, bias detection, empathy coaching, inclusion monitoring, recognition systems, mood analysis, conflict mediation — builds healthy culture'),
    ('Governance & Logic (8)', 'POSH compliance, labor law, data privacy, audit trails, conflict-of-interest detection, leave optimization, policy interpretation — ensures compliance'),
]
for name, desc in clusters:
    add_bullet(desc, bold_prefix=name)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Key differentiator: ')
run.bold = True
p.add_run('Most HR platforms have 1 chatbot. We have 70 specialized agents working as a team, '
          'like a hospital with specialists vs a single general doctor.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 7. RBAC
# ═══════════════════════════════════════════════════════════
doc.add_heading('7. Who Sees What — Roles & Access Control (RBAC)', level=1)

doc.add_paragraph(
    'RBAC stands for Role-Based Access Control. It ensures that people only see '
    'what they are supposed to see. A regular employee cannot see another employee\'s salary or review.'
)

doc.add_heading('The 5 Role Levels', level=2)

role_table = doc.add_table(rows=1, cols=4)
role_table.style = 'Light Grid Accent 1'
hdr = role_table.rows[0].cells
for i, text in enumerate(['Role', 'Who Is This?', 'What Can They See?', 'What Can They Do?']):
    hdr[i].text = text

roles = [
    ('SUPER_ADMIN', 'Platform Owner', 'Everything across ALL organizations', 'Create tenants, manage subscriptions, system settings, view all data'),
    ('ADMIN', 'Tenant Admin (Org Owner)', 'Everything within their own organization', 'Manage users, roles, licenses, settings, audit logs for their org'),
    ('MANAGER', 'Team Manager / HR Admin', 'Their direct reports + team data', 'Approve goals, conduct reviews, team analytics, calibration'),
    ('HR_BP', 'HR Business Partner', 'Department-level employee data', 'HR-specific admin functions, engagement analysis, compliance'),
    ('EMPLOYEE', 'Regular Employee', 'Only their own data', 'Set goals, submit self-reviews, give feedback, view own scores'),
]
for r in roles:
    add_table_row(role_table, r, bold_first=True)

doc.add_paragraph()
doc.add_heading('How Access Scope Works', level=2)

scopes = [
    ('"Own"', 'Can only see your own data (employees)'),
    ('"Team"', 'Can see direct reports (managers)'),
    ('"Department"', 'Can see everyone in your department (department heads)'),
    ('"Business Unit"', 'Can see everyone in your business unit (BU heads)'),
    ('"All"', 'Can see everyone in the organization (admins)'),
]
for scope, desc in scopes:
    add_bullet(desc, bold_prefix=scope)

doc.add_paragraph()
doc.add_heading('Security Features', level=2)
add_bullet('Cross-tenant access is blocked and logged as a security event')
add_bullet('Brute force detection: 5+ failed logins in an hour triggers an alert')
add_bullet('Bulk deactivation alerts: 5+ deactivations in an hour triggers investigation')
add_bullet('Every action is recorded in an immutable audit log')
add_bullet('Delegation of authority — a manager can temporarily grant access to another user')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 8. SUPER ADMIN vs TENANT ADMIN
# ═══════════════════════════════════════════════════════════
doc.add_heading('8. Super Admin vs Tenant Admin', level=1)

doc.add_paragraph(
    'Think of it like this: Super Admin is the LANDLORD of the building. '
    'Tenant Admin is the OWNER of an apartment in that building.'
)

sa_table = doc.add_table(rows=1, cols=3)
sa_table.style = 'Light Grid Accent 1'
hdr = sa_table.rows[0].cells
for i, text in enumerate(['Aspect', 'Super Admin (Landlord)', 'Tenant Admin (Apartment Owner)']):
    hdr[i].text = text

comparisons = [
    ('Scope', 'Entire platform, all organizations', 'Only their own organization'),
    ('App', 'Separate Admin Portal (/admin/)', 'Main Web App (/web/)'),
    ('Create Organizations?', 'Yes — can create new tenants', 'No'),
    ('Manage Users?', 'All users across all orgs', 'Only users in their org'),
    ('Manage Subscriptions?', 'Yes — billing, upgrades, suspensions', 'View-only (request upgrades)'),
    ('See Other Orgs?', 'Yes (with audit logging)', 'Never — completely isolated'),
    ('AI Features?', 'Always available', 'Only if subscription includes it'),
    ('System Settings?', 'Global platform settings', 'Org-level settings only'),
    ('Audit Logs?', 'All platform events', 'Own org events only'),
    ('License Management?', 'View all tenant licenses', 'Manage own tenant seats'),
]
for c in comparisons:
    add_table_row(sa_table, c, bold_first=True)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 9. MULTI-TENANT
# ═══════════════════════════════════════════════════════════
doc.add_heading('9. Multi-Tenant SaaS Architecture (Simple)', level=1)

doc.add_heading('What Is Multi-Tenant?', level=2)
doc.add_paragraph(
    'One software installation serves many organizations (tenants). '
    'Like Gmail — millions of companies use the same Gmail, but Company A '
    'can never see Company B\'s emails. Each company gets its own isolated space.'
)

doc.add_heading('How We Do It', level=2)
add_bullet('Every piece of data has an organization tag (tenantId)')
add_bullet('Every database query automatically filters by organization')
add_bullet('Authorization middleware blocks cross-organization access')
add_bullet('Security events are logged if someone tries to access another org')
add_bullet('Each organization has its own: users, settings, subscription, branding')

doc.add_heading('Each Tenant Gets', level=2)
add_bullet('A unique slug (e.g., acme-corp) for login')
add_bullet('Their own user directory and org structure')
add_bullet('Custom branding and settings')
add_bullet('A designated manager for data uploads')
add_bullet('Their own subscription plan and license count')
add_bullet('Isolated audit logs and security events')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 10. SUBSCRIPTION PLANS
# ═══════════════════════════════════════════════════════════
doc.add_heading('10. Subscription Plans & Feature Gating', level=1)

plan_table = doc.add_table(rows=1, cols=5)
plan_table.style = 'Light Grid Accent 1'
hdr = plan_table.rows[0].cells
for i, text in enumerate(['Feature', 'FREE', 'STARTER', 'PROFESSIONAL', 'ENTERPRISE']):
    hdr[i].text = text

Y = 'Yes'
N = '—'
plans = [
    ('Basic Goals & Reviews', Y, Y, Y, Y),
    ('Feedback & Recognition', Y, Y, Y, Y),
    ('My Performance Dashboard', Y, Y, Y, Y),
    ('Team Overview Dashboard', Y, Y, Y, Y),
    ('HR Analytics Dashboard', N, Y, Y, Y),
    ('Organization Dashboard', N, Y, Y, Y),
    ('Engagement Heatmap', N, Y, Y, Y),
    ('Feedback Analysis', N, Y, Y, Y),
    ('Calibration Sessions', N, N, Y, Y),
    ('Bias Detection', N, N, Y, Y),
    ('Attrition Risk', N, N, Y, Y),
    ('70 AI Agents', N, N, Y, Y),
    ('Advanced Analytics', N, N, Y, Y),
    ('Compensation Management', N, N, N, Y),
    ('Custom Integrations', N, N, N, Y),
    ('Dedicated Support', N, N, N, Y),
]
for p_data in plans:
    add_table_row(plan_table, p_data, bold_first=True)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 11. REAL-TIME
# ═══════════════════════════════════════════════════════════
doc.add_heading('11. Real-Time & Live Updates', level=1)

doc.add_paragraph(
    'The dashboard shows a green pulsing "Live" indicator when connected. '
    'Data updates automatically without refreshing the page.'
)

doc.add_heading('How It Works', level=2)
add_bullet('WebSocket connection: Like a phone call that stays open — server pushes updates instantly')
add_bullet('When data changes (e.g., someone submits a review), the server notifies all connected users')
add_bullet('Charts and scores refresh automatically within seconds')
add_bullet('If the connection drops, the system falls back to polling (checking every 30 seconds)')
add_bullet('Visual indicator shows connection status: Green = Live, Amber = Connecting, Red = Offline')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 12. SECURITY
# ═══════════════════════════════════════════════════════════
doc.add_heading('12. Security & Compliance', level=1)

security_items = [
    ('Data Isolation', 'Every organization\'s data is completely separate. Even our own engineers can\'t see one company\'s data from another.'),
    ('Audit Logging', 'Every action — login, data access, changes — is permanently recorded with who, what, when, and where.'),
    ('Brute Force Protection', 'After 5 failed login attempts in an hour, the system alerts admins and can lock the account.'),
    ('Cross-Tenant Blocking', 'If someone tries to access another organization\'s data (even through API hacking), it\'s blocked and logged.'),
    ('Role-Based Access', '5 levels of access control ensure employees only see what they need to see.'),
    ('Bias Detection', 'AI continuously monitors for rating bias, discrimination patterns, and fairness violations.'),
    ('Delegation Controls', 'Temporary access grants are tracked and expire automatically.'),
    ('Encryption', 'All data encrypted in transit (HTTPS/TLS) and at rest in the database.'),
]
for name, desc in security_items:
    add_bullet(desc, bold_prefix=name)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 13. KEY NUMBERS
# ═══════════════════════════════════════════════════════════
doc.add_heading('13. Key Numbers to Remember', level=1)

numbers_table = doc.add_table(rows=1, cols=2)
numbers_table.style = 'Light Grid Accent 1'
hdr = numbers_table.rows[0].cells
hdr[0].text = 'Metric'
hdr[1].text = 'Number'

numbers = [
    ('Total Features', '130+'),
    ('Web Pages', '87+'),
    ('API Modules', '44'),
    ('Database Models', '180+'),
    ('AI Agents', '70'),
    ('Agent Clusters', '6'),
    ('CPIS Dimensions', '8'),
    ('Statistical Formulas', '17+'),
    ('User Roles', '5 levels + custom roles'),
    ('Subscription Plans', '4 (Free, Starter, Professional, Enterprise)'),
    ('Dashboard Charts', '29 across 4 tabs'),
    ('Review Types', '5 (Self, Manager, Peer, 360, External)'),
    ('Real-Time Update', 'WebSocket + 30s polling fallback'),
    ('PWA Support', 'Yes (works offline on mobile)'),
    ('Dark Mode', 'Full support'),
]
for n in numbers:
    add_table_row(numbers_table, n, bold_first=True)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 14. DEMO SCRIPT
# ═══════════════════════════════════════════════════════════
doc.add_heading('14. Demo Flow — Recommended 5-Minute Script', level=1)

script = [
    ('0:00 - 0:30 | Login & First Impression',
     'Log in as an Employee. Show the login screen with organization ID. '
     'Point out the premium glassmorphic design and dark mode support.'),
    ('0:30 - 1:30 | My Performance Dashboard',
     'Show the CPIS score with the animated ring gauge and 8-dimension radar chart. '
     'Scroll down to show the 12 charts: goal progress, skill gap radar, performance timeline, '
     'activity heatmap, and feedback sentiment. Mention the scroll animations and live indicator.'),
    ('1:30 - 2:30 | AI Agent Demo',
     'Open the AI workspace. Ask "Draft my quarterly self-review." '
     'Show how the AI reads goals, feedback, and achievements, then generates a complete review. '
     'Mention that it\'s one of 70 specialized agents.'),
    ('2:30 - 3:30 | Manager View',
     'Switch to Manager login. Show the Team Overview tab with 8 charts. '
     'Point out team health gauge, performance distribution, review cycle progress. '
     'Show how managers can drill down into individual employee profiles.'),
    ('3:30 - 4:15 | HR Analytics & Admin',
     'Switch to HR Admin login. Show the HR Analytics tab with engagement heatmap, '
     'bias detection, attrition risk scatter chart. Mention calibration sessions.'),
    ('4:15 - 5:00 | Closing Highlights',
     'Mention: 70 AI agents, 17+ math formulas, real-time WebSocket updates, '
     'multi-tenant SaaS, enterprise security, PWA support, 130+ features. '
     'Show the subscription plan gating (blurred charts for lower plans).'),
]

for step_title, step_desc in script:
    p = doc.add_paragraph()
    run = p.add_run(step_title)
    run.bold = True
    run.font.color.rgb = RGBColor(0x6D, 0x28, 0xD9)
    doc.add_paragraph(step_desc)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 15. FAQs
# ═══════════════════════════════════════════════════════════
doc.add_heading('15. 30 FAQs with Quick Answers', level=1)

faqs = [
    ('What is PMS Platform?', 'An AI-powered SaaS platform that manages, measures, and improves employee performance across organizations.'),
    ('What does CPIS stand for?', 'Comprehensive Performance Intelligence Score — a single 0-100 number that captures 8 dimensions of performance.'),
    ('How is the CPIS score calculated?', 'It uses 17+ statistical formulas across 8 weighted dimensions: goals (25%), reviews (20%), feedback (12%), collaboration (10%), consistency (10%), growth (8%), evidence (8%), initiative (7%).'),
    ('What makes this different from other PMS tools?', 'Three things: 70 specialized AI agents (not one chatbot), mathematically rigorous scoring with bias detection, and full multi-tenant SaaS architecture.'),
    ('What is Agentic AI?', 'AI that can take actions — not just answer questions, but draft reviews, create goals, generate reports, and schedule meetings autonomously.'),
    ('How many AI agents does the platform have?', '70 specialized agents across 6 clusters: Core, Bio-Performance, Hyper-Learning, Liquid Workforce, Culture & Empathy, and Governance.'),
    ('Is employee data safe across organizations?', 'Yes — every organization is completely isolated. Cross-tenant access is blocked and logged as a security event.'),
    ('What is multi-tenant architecture?', 'One software serves many organizations, each with completely isolated data, like how Gmail works for different companies.'),
    ('What roles does the system support?', 'Five built-in roles: Super Admin, Tenant Admin, Manager, HR Business Partner, and Employee — plus custom roles.'),
    ('What is the difference between Super Admin and Tenant Admin?', 'Super Admin manages the entire platform (like a landlord), Tenant Admin manages one organization (like an apartment owner).'),
    ('How does bias detection work?', 'The system uses Z-score normalization to detect reviewer bias, the 4/5ths rule for disparate impact, and flags self-review inflation.'),
    ('What is calibration?', 'A process where HR normalizes performance ratings across teams so one strict manager doesn\'t disadvantage their team.'),
    ('Does it work on mobile?', 'Yes — it\'s a Progressive Web App (PWA) that works on any device without app store installation.'),
    ('How do real-time updates work?', 'WebSocket connections push data changes instantly. If the connection drops, it falls back to polling every 30 seconds.'),
    ('What subscription plans are available?', 'Four tiers: Free (basic features), Starter (+ HR analytics), Professional (+ AI agents, bias detection), Enterprise (all features).'),
    ('How many features does the platform have?', 'Over 130 distinct features across 11 categories, with 87+ web pages and 44 API modules.'),
    ('What is a review cycle?', 'A structured period (quarterly or annual) where employees receive performance reviews from managers, peers, and self-assessments.'),
    ('What is 360-degree feedback?', 'Feedback collected from all directions — managers, peers, direct reports, self, and external reviewers — for a complete picture.'),
    ('How does the scoring avoid penalizing new employees?', 'Bayesian estimation: when data is limited, the system borrows from department averages and gradually shifts to individual data as more evidence is collected.'),
    ('What is eNPS?', 'Employee Net Promoter Score — measures how likely employees are to recommend the company as a great place to work, on a -100 to +100 scale.'),
    ('Can managers see other departments\' data?', 'No — managers only see their direct reports and team. Department heads see their department only. Cross-access is blocked.'),
    ('What happens if someone tries to access another organization\'s data?', 'The system blocks the request, logs it as a security event, and alerts admins if repeated attempts are detected.'),
    ('What is the goal hierarchy?', 'Goals can be linked parent-to-child, connecting individual goals to team goals to company objectives (OKR alignment).'),
    ('How does the skill gap analysis work?', 'It compares current skill levels (rated 1-5) against target levels set by managers, and shows the gap visually on a radar chart.'),
    ('What is a PIP?', 'Performance Improvement Plan — a structured program for underperforming employees with clear milestones, regular check-ins, and documented outcomes.'),
    ('How does the leaderboard work?', 'Employees are ranked by a composite of goal completion, review ratings, and feedback sentiment — viewable weekly, monthly, or yearly.'),
    ('What is the audit trail?', 'A permanent, tamper-proof record of every action taken in the system — who did what, when, and from where.'),
    ('Can custom roles be created?', 'Yes — Tenant Admins can create custom roles with specific permissions tailored to their organization\'s structure.'),
    ('How does delegation work?', 'A manager can temporarily grant their authority to another user (e.g., during vacation). It\'s tracked and expires automatically.'),
    ('Is the platform legally defensible?', 'Yes — the mathematical scoring, bias detection, audit trails, and calibration processes create a documentation trail that meets legal standards.'),
]

for i, (q, a) in enumerate(faqs, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'Q{i}: {q}')
    run.bold = True
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    p = doc.add_paragraph(f'A: {a}')
    p.paragraph_format.space_after = Pt(8)

# ═══════════════════════════════════════════════════════════
# BACK PAGE
# ═══════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
closing = doc.add_paragraph()
closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = closing.add_run('PMS Platform')
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x6D, 0x28, 0xD9)
run.bold = True

tagline2 = doc.add_paragraph()
tagline2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = tagline2.add_run('Intelligent Performance. Measurable Growth. Fair for Everyone.')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
run.italic = True

doc.add_paragraph()
contact = doc.add_paragraph()
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = contact.add_run('pms.xzashr.com  |  Confidential  |  March 2026')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

# ── Save ──────────────────────────────────────────────────
output_path = os.path.join(os.path.dirname(__file__), 'PMS_Demo_Preparation_Guide.docx')
doc.save(output_path)
print(f'Document saved to: {output_path}')
