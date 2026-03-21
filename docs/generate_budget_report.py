"""
Generate PMS AI Budget & Model Report (.docx)
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# Styles
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

for lvl in range(1, 4):
    h = doc.styles[f'Heading {lvl}']
    h.font.name = 'Calibri'
h1s = doc.styles['Heading 1']
h1s.font.color.rgb = RGBColor(0x6D, 0x28, 0xD9)
h1s.font.size = Pt(18)
h2s = doc.styles['Heading 2']
h2s.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
h2s.font.size = Pt(14)

VIOLET = RGBColor(0x6D, 0x28, 0xD9)
BLUE = RGBColor(0x1E, 0x3A, 0x5F)
GRAY = RGBColor(0x64, 0x74, 0x8B)
NAVY = RGBColor(0x0F, 0x17, 0x2A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK_HEADER = "1E293B"
ALT_ROW = "F1F5F9"

def shade(cell, color):
    tc = cell._element.get_or_add_tcPr()
    shd = tc.makeelement(qn('w:shd'), {qn('w:fill'): color, qn('w:val'): 'clear'})
    tc.append(shd)

def make_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    hdr = t.rows[0]
    for i, txt in enumerate(headers):
        c = hdr.cells[i]
        shade(c, DARK_HEADER)
        p = c.paragraphs[0]
        r = p.add_run(txt)
        r.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = WHITE
        r.font.name = 'Calibri'
    # Data rows
    for ri, row_data in enumerate(rows):
        row = t.add_row()
        for ci, txt in enumerate(row_data):
            c = row.cells[ci]
            if ri % 2 == 1:
                shade(c, ALT_ROW)
            p = c.paragraphs[0]
            r = p.add_run(str(txt))
            r.font.size = Pt(9.5)
            r.font.name = 'Calibri'
            if ci == 0:
                r.bold = True
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    return t

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(' ' + text)
    else:
        p.add_run(text)

def note(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(9.5)
    r.font.color.rgb = GRAY
    r.italic = True

# ══════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('PMS Platform')
r.font.size = Pt(36)
r.font.color.rgb = VIOLET
r.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Agentic AI Budget & Model Report')
r.font.size = Pt(22)
r.font.color.rgb = BLUE

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Token Usage  |  Model Comparison  |  Cost Estimation  |  Recommendations')
r.font.size = Pt(11)
r.font.color.rgb = GRAY

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Prepared for: Management & Stakeholders')
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('March 2026  |  Confidential')
r.font.size = Pt(10)
r.font.color.rgb = GRAY
r.italic = True

doc.add_page_break()

# ══════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ══════════════════════════════════════
doc.add_heading('1. Executive Summary', level=1)
doc.add_paragraph(
    'This report outlines the token usage, model options, cost projections, and recommendations '
    'for the PMS Platform\u2019s Agentic AI system. The platform features a sophisticated '
    'multi-provider AI infrastructure designed for enterprise-grade performance management.'
)
make_table(
    ['Metric', 'Value'],
    [
        ['AI Agents', '70 specialized agents across 6 clusters'],
        ['Tools Available', '64 registered tools (27 read + 37 write)'],
        ['LLM Providers', '5 (Anthropic, OpenAI, Google Gemini, DeepSeek, Groq)'],
        ['Primary Model', 'Claude Sonnet 4 (claude-sonnet-4-20250514)'],
        ['Default Max Tokens', '4,096 per response'],
        ['Per-Task Token Cap', '50,000 tokens (cumulative)'],
        ['Per-Task Cost Cap', '$0.50 (circuit breaker)'],
        ['Response Caching', 'Redis with 1-hour TTL'],
        ['Rate Limits', '15 calls/user/min, 60 calls/tenant/min'],
    ],
    [2.5, 4.5]
)

doc.add_page_break()

# ══════════════════════════════════════
# 2. TOKEN USAGE
# ══════════════════════════════════════
doc.add_heading('2. Token Usage Per Interaction Type', level=1)
doc.add_paragraph(
    'Each user interaction triggers a different number of LLM calls depending on complexity. '
    'Below are the token ranges and costs using the primary model (Claude Sonnet 4).'
)
make_table(
    ['Interaction Type', 'LLM Calls', 'Input Tokens', 'Output Tokens', 'Cost / Interaction'],
    [
        ['Simple Chat (Q&A)', '1', '2,000 \u2013 4,000', '500 \u2013 1,500', '$0.01 \u2013 $0.03'],
        ['Agentic Task (multi-step)', '3 \u2013 5', '10,000 \u2013 20,000', '5,000 \u2013 10,000', '$0.05 \u2013 $0.20'],
        ['Help Assistant', '1', '1,500 \u2013 2,500', '300 \u2013 1,024', '$0.005 \u2013 $0.02'],
        ['Proactive System Scan', '1 \u2013 2', '1,000 \u2013 3,000', '500 \u2013 1,500', '$0.01 \u2013 $0.05'],
    ],
    [1.6, 0.9, 1.3, 1.3, 1.5]
)

doc.add_paragraph()
doc.add_heading('How Agentic Tasks Use Tokens (Breakdown)', level=2)
bullet('Intent classification: ~100 tokens', 'Step 1 \u2014')
bullet('Task planning: up to 4,096 tokens', 'Step 2 \u2014')
bullet('Tool execution + observation: ~2,048 tokens per tool call (up to 15 steps)', 'Step 3 \u2014')
bullet('Recovery on failure: ~1,024 tokens (if needed)', 'Step 4 \u2014')
bullet('Task summary: ~1,500 tokens', 'Step 5 \u2014')
doc.add_paragraph('Maximum 15 steps per agentic task. Cumulative cap: 50,000 tokens or $0.50 per task.')

doc.add_page_break()

# ══════════════════════════════════════
# 3. MODEL COMPARISON
# ══════════════════════════════════════
doc.add_heading('3. Model Comparison & Pricing', level=1)
doc.add_paragraph('The platform supports 5 LLM providers with automatic fallback:')
make_table(
    ['Model', 'Provider', 'Input/1K', 'Output/1K', 'Quality', 'Speed', 'Best For'],
    [
        ['Claude Sonnet 4', 'Anthropic', '$0.003', '$0.015', 'Excellent', 'Medium', 'Production (primary)'],
        ['GPT-4o', 'OpenAI', '$0.0025', '$0.010', 'Excellent', 'Fast', 'Fallback provider'],
        ['GPT-4o Mini', 'OpenAI', '$0.00015', '$0.0006', 'Good', 'Very Fast', 'Economy testing'],
        ['Gemini 2.0 Flash', 'Google', '$0.0001', '$0.0004', 'Good', 'Very Fast', 'Economy tier'],
        ['DeepSeek Chat', 'DeepSeek', '$0.00014', '$0.00028', 'Good', 'Fast', 'Budget alternative'],
        ['Groq Llama 3.3 70B', 'Groq', '$0.00059', '$0.00079', 'Good', 'Ultra Fast', 'Speed-critical'],
        ['Claude 3 Haiku', 'Anthropic', '$0.00025', '$0.00125', 'Moderate', 'Very Fast', 'Help Assistant'],
    ],
    [1.4, 0.8, 0.7, 0.7, 0.7, 0.8, 1.2]
)
note('Gemini 2.0 Flash is ~30x cheaper than Claude Sonnet 4 for input and ~37x cheaper for output.')

doc.add_page_break()

# ══════════════════════════════════════
# 4. MONTHLY COST ESTIMATION
# ══════════════════════════════════════
doc.add_heading('4. Monthly Cost Estimation by User Count', level=1)
doc.add_paragraph(
    'Assumptions: 10 chat interactions/user/day + 2 agentic tasks/user/day + 5 system scans/day. '
    '22 working days/month. Claude Sonnet 4 pricing.'
)
make_table(
    ['Users', 'Chat Cost/mo', 'Agentic Cost/mo', 'System/mo', 'Total / Month'],
    [
        ['50', '$110 \u2013 $330', '$110 \u2013 $440', '$2 \u2013 $12', '$220 \u2013 $780'],
        ['100', '$220 \u2013 $660', '$220 \u2013 $880', '$2 \u2013 $12', '$440 \u2013 $1,550'],
        ['250', '$550 \u2013 $1,650', '$550 \u2013 $2,200', '$5 \u2013 $24', '$1,100 \u2013 $3,870'],
        ['500', '$1,100 \u2013 $3,300', '$1,100 \u2013 $4,400', '$10 \u2013 $48', '$2,200 \u2013 $7,750'],
        ['1,000', '$2,200 \u2013 $6,600', '$2,200 \u2013 $8,800', '$10 \u2013 $48', '$4,400 \u2013 $15,450'],
    ],
    [0.8, 1.5, 1.5, 1.4, 1.5]
)
note('Upper-bound estimates. Actual costs 40-60% lower with caching, economy routing, and rate limiting.')

doc.add_paragraph()
doc.add_heading('Optimized Cost (with Economy Tier Routing)', level=2)
doc.add_paragraph('Route simple queries to Gemini 2.0 Flash, reserve Claude Sonnet 4 for complex agentic tasks:')
make_table(
    ['Users', 'Optimized Total / Month', 'Savings vs Full Claude'],
    [
        ['50', '$90 \u2013 $320', '~55% savings'],
        ['100', '$180 \u2013 $640', '~55% savings'],
        ['250', '$450 \u2013 $1,600', '~55% savings'],
        ['500', '$900 \u2013 $3,200', '~55% savings'],
        ['1,000', '$1,800 \u2013 $6,400', '~55% savings'],
    ],
    [2.3, 2.3, 2.3]
)

doc.add_page_break()

# ══════════════════════════════════════
# 5. TESTING RECOMMENDATION
# ══════════════════════════════════════
doc.add_heading('5. Recommendation: Testing Phase', level=1)
make_table(
    ['Aspect', 'Recommendation'],
    [
        ['Model', 'GPT-4o Mini or Gemini 2.0 Flash'],
        ['Why', '50\u2013100x cheaper than Claude Sonnet 4'],
        ['Cost (50 users, 1 month)', '$5 \u2013 $15'],
        ['Cost (100 users, 1 month)', '$10 \u2013 $30'],
        ['Quality', 'Sufficient for UI testing, basic flows, demos'],
        ['Limitations', 'Lower reasoning; may fail complex multi-step agentic tasks'],
        ['Configuration', 'Set AI_PRIMARY_PROVIDER=gemini or openai in .env'],
    ],
    [2.5, 4.5]
)
doc.add_paragraph()
doc.add_paragraph('For initial QA and demos, economy models are sufficient. Switch to Claude Sonnet 4 only when validating agentic task quality (multi-step workflows, review drafting, analytics).')

# ══════════════════════════════════════
# 6. PRODUCTION RECOMMENDATION
# ══════════════════════════════════════
doc.add_heading('6. Recommendation: Production Deployment', level=1)
make_table(
    ['Tier', 'Model', 'Use Case', 'Cost Level'],
    [
        ['Primary', 'Claude Sonnet 4', 'Agentic tasks, review drafting, complex analytics', '$$$$'],
        ['Economy', 'Gemini 2.0 Flash', 'Help Assistant, simple Q&A, cached queries', '$'],
        ['Fallback', 'GPT-4o', 'Auto-failover when Anthropic is unavailable', '$$$'],
    ],
    [1.0, 1.7, 2.8, 1.0]
)

doc.add_paragraph()
doc.add_heading('Built-in Cost Optimization', level=2)
bullet('Redis caching (1hr TTL) \u2014 identical questions cost zero on repeat', 'Caching:')
bullet('15 calls/user/min, 60/tenant/min \u2014 prevents cost spikes', 'Rate Limiting:')
bullet('Trips after 3 failures, 5-min cooldown, auto-fallback', 'Circuit Breaker:')
bullet('50K tokens/task, $0.50/task \u2014 no runaway costs', 'Token Caps:')
bullet('Configurable via AI_MONTHLY_BUDGET_CENTS (default $100/tenant)', 'Monthly Budget:')

doc.add_page_break()

# ══════════════════════════════════════
# 7. SAFEGUARDS
# ══════════════════════════════════════
doc.add_heading('7. Built-in Safeguards', level=1)
doc.add_paragraph('Comprehensive cost and safety controls prevent overspending:')
make_table(
    ['Safeguard', 'Limit', 'When Triggered'],
    [
        ['Per-Task Token Limit', '50,000 tokens', 'Task terminates with summary'],
        ['Per-Task Cost Limit', '$0.50 max', 'Task terminates, user notified'],
        ['Per-User Rate Limit', '15 calls/minute', '429 error with exponential backoff'],
        ['Per-Tenant Rate Limit', '60 calls/minute', 'Tenant-wide throttle'],
        ['Monthly Budget', 'Default $100', 'AI pauses until next billing cycle'],
        ['Circuit Breaker', '3 consecutive failures', '5-min cooldown, auto-switches provider'],
        ['Response Cache', '1-hour TTL (Redis)', 'Cached response returned, zero LLM cost'],
        ['Provider Fallback', '5 providers in chain', 'Anthropic \u2192 OpenAI \u2192 Gemini \u2192 DeepSeek \u2192 Groq'],
        ['Agentic Step Limit', '15 steps per task', 'Task terminates with partial results'],
        ['Human-in-the-Loop', 'Write operations', 'User must approve risky AI actions'],
    ],
    [2.0, 2.0, 3.0]
)

# ══════════════════════════════════════
# 8. API KEYS
# ══════════════════════════════════════
doc.add_heading('8. API Key Requirements', level=1)
doc.add_paragraph('The following API keys need to be configured in the environment (.env file):')
make_table(
    ['Provider', 'Env Variable', 'Required?', 'Notes'],
    [
        ['Anthropic', 'ANTHROPIC_API_KEY', 'Yes (Primary)', 'Required for production'],
        ['OpenAI', 'OPENAI_API_KEY', 'Recommended', 'Fallback when Anthropic down'],
        ['Google', 'GEMINI_API_KEY', 'Optional', 'Economy tier (30x cheaper)'],
        ['DeepSeek', 'DEEPSEEK_API_KEY', 'Optional', 'Budget alternative'],
        ['Groq', 'GROQ_API_KEY', 'Optional', 'Ultra-fast inference'],
    ],
    [1.2, 2.0, 1.3, 2.5]
)
note('Minimum: Only ANTHROPIC_API_KEY is needed. Additional keys enable fallback and cost optimization.')

doc.add_page_break()

# ══════════════════════════════════════
# 9. DECISION MATRIX
# ══════════════════════════════════════
doc.add_heading('9. Quick Decision Matrix', level=1)
doc.add_paragraph('Use this table to quickly pick the right model and budget for your scenario:')
make_table(
    ['Scenario', 'Model', 'Cost (100 users/mo)', 'Notes'],
    [
        ['Testing / Demo', 'Gemini 2.0 Flash', '$8 \u2013 $25', 'Cheapest, good for UI testing'],
        ['Pilot (limited AI)', 'GPT-4o Mini', '$15 \u2013 $50', 'Better quality, affordable'],
        ['Production (standard)', 'Claude Sonnet 4', '$440 \u2013 $1,550', 'Full quality, all features'],
        ['Production (optimized)', 'Claude + Gemini', '$180 \u2013 $640', 'Best value: premium + economy'],
        ['Enterprise (1000+)', 'Claude + GPT-4o', '$1,800 \u2013 $6,400', 'Optimized with fallback'],
    ],
    [1.7, 1.7, 1.7, 2.0]
)

# ══════════════════════════════════════
# 10. CONCLUSION
# ══════════════════════════════════════
doc.add_paragraph()
doc.add_heading('10. Conclusion', level=1)
doc.add_paragraph('The PMS Platform\u2019s Agentic AI system is architected for cost-efficiency and reliability:')
bullet('5 providers with auto-fallback ensures zero AI downtime.', 'Resilience \u2014')
bullet('Caching, economy routing, and rate limiting reduce costs by 40-60%.', 'Cost Control \u2014')
bullet('Token caps, cost caps, and monthly budgets prevent overspending.', 'Safeguards \u2014')
bullet('Switch models by changing one env variable. No code changes needed.', 'Flexibility \u2014')
bullet('Start at $8/month (testing), scale to $640/month (100 users, optimized).', 'Scalability \u2014')

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Recommended: Start with Gemini 2.0 Flash for testing ($8\u201325/mo),')
r.bold = True
r.font.color.rgb = VIOLET
r.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('deploy with Claude Sonnet 4 + Gemini economy ($180\u2013640/mo for 100 users).')
r.bold = True
r.font.color.rgb = VIOLET
r.font.size = Pt(12)

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('PMS Platform  |  Confidential  |  March 2026')
r.font.size = Pt(10)
r.font.color.rgb = GRAY
r.italic = True

# Save
output_path = os.path.join(os.path.dirname(__file__), 'PMS_AI_Budget_Report.docx')
doc.save(output_path)
print(f'Saved: {output_path}')
